# utils/doc_utils.py
import io
from docx import Document

def extract_text_from_docx_stream(file_stream):
    """
    file_stream: BytesIO or similar
    """
    file_stream.seek(0)
    doc = Document(file_stream)
    text = "\n".join([p.text for p in doc.paragraphs])
    return text

def extract_text_docx_and_strip_metadata(filepath):
    """
    If you have a saved .docx on disk and want to strip metadata:
    """
    doc = Document(filepath)
    core = doc.core_properties
    core.author = ""
    core.comments = ""
    core.title = ""
    core.subject = ""
    core.keywords = ""
    doc.save(filepath)
